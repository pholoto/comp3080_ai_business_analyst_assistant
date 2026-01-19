"""Generate BA reports using the provided DOCX template."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..config import WORKSPACE_ROOT
from ..memory import Session


class BAReportGenerator:
    """Render a structured report aligned with the VinUni CECS capstone template."""

    def __init__(self, templates_dir: Path | None = None) -> None:
        self.templates_dir = templates_dir or (WORKSPACE_ROOT / "back_end" / "templates")

    def generate(self, session: Session, output_path: Path | None = None, report_type: str = "BA Report") -> Path:
        """Create the report populated with the session artefacts."""
        if output_path is None:
            output_dir_str = session.state.get("report_output_dir")
            if output_dir_str:
                output_dir = Path(output_dir_str)
            else:
                # Fallback to shared data dir if not set
                output_dir = WORKSPACE_ROOT / "back_end" / "data" / "shared"
            
            output_dir.mkdir(parents=True, exist_ok=True)
            filename = f"{report_type.replace(' ', '_')}_{session.session_id}.docx"
            output_path = output_dir / filename
            
        # Try to find a matching template
        template_path = self._find_template(report_type)
        if template_path:
            document = Document(str(template_path))
        else:
            document = Document() # Create blank document if no template found

        self._apply_heading_style(document)
        
        if report_type == "SRS":
            self._populate_srs(document, session)
        elif report_type == "BA Report":
            self._populate_document(document, session)
        else:
            # Default to full report population for unknown types for now
            # In a future iteration, this could be a generic dump
            document.add_heading(report_type, level=0)
            self._populate_document(document, session)
            
        document.save(str(output_path))
        session.set_state("report_path", str(output_path))
        return output_path

    def _find_template(self, report_type: str) -> Path | None:
        """Search for a template matching the report type."""
        candidates = [
            self.templates_dir / f"{report_type}.docx",
            self.templates_dir / f"{report_type.replace(' ', '_')}.docx",
            self.templates_dir / f"{report_type.replace(' ', '-')}.docx",
        ]
        # Legacy default for BA Report
        if report_type == "BA Report":
            candidates.append(self.templates_dir / "VinUni-CECS-Capstone-Project-template.docx")
            
        for cand in candidates:
            if cand.exists():
                return cand
        return None

    def _apply_heading_style(self, document: Document) -> None:
        """Ensure heading style is consistent even if template is missing definitions."""
        try:
            style = document.styles["Heading 1"]
            if style.font:
                style.font.name = "Calibri"
                style.font.size = Pt(16)
        except KeyError:
            pass

    def _populate_document(self, document: Document, session: Session) -> None:
        self._write_project_overview(document, session)
        self._write_problem_definition(document, session)
        self._write_requirements(document, session)
        self._write_solution_design(document, session)
        self._write_prototype_development(document, session)
        self._write_testing_validation(document, session)
        self._write_documentation_summary(document, session)
        self._write_market_analysis(document, session)
        self._write_attachments(document, session)
        self._write_conversation_log(document, session)

    def _populate_srs(self, document: Document, session: Session) -> None:
        document.add_heading("Software Requirements Specification", level=1)
        self._write_project_overview(document, session)
        self._write_problem_definition(document, session)
        self._write_requirements(document, session)
        self._write_solution_design(document, session)
        self._write_testing_validation(document, session)
        self._write_attachments(document, session)

    def _write_project_overview(self, document: Document, session: Session) -> None:
        document.add_heading("Project Overview", level=1)
        overview = session.get_state("project_overview") or "Project overview not captured yet."
        document.add_paragraph(overview)

    def _write_problem_definition(self, document: Document, session: Session) -> None:
        document.add_heading("Problem Definition", level=1)
        problem_statement = session.get_state("problem_statement") or "Problem statement not captured."
        document.add_paragraph(problem_statement)
        self._write_simple_list(document, session.get_state("pain_points"), "Pain Points")
        self._write_simple_list(document, session.get_state("success_metrics"), "Success Metrics")
        self._write_simple_list(document, session.get_state("clarifying_questions"), "Clarifying Questions")
        self._write_simple_list(document, session.get_state("assumptions"), "Assumptions")

    def _write_requirements(self, document: Document, session: Session) -> None:
        document.add_heading("Requirements Analysis", level=1)
        requirements = session.get_state("requirements") or []
        if requirements:
            document.add_paragraph("Functional Requirements:")
            for item in requirements:
                requirement = item if isinstance(item, Mapping) else {"requirement": str(item)}
                req_id = requirement.get("id")
                req_text = requirement.get("requirement") or requirement.get("name") or str(item)
                priority = requirement.get("priority")
                rationale = requirement.get("rationale")
                label = f"• {req_id}: {req_text}" if req_id else f"• {req_text}"
                if priority:
                    label += f" (Priority: {priority})"
                document.add_paragraph(label)
                if rationale:
                    document.add_paragraph(f"  Rationale: {rationale}")
        else:
            document.add_paragraph("Functional requirements have not been documented.")
        non_functional = session.get_state("non_functional_requirements") or []
        if non_functional:
            document.add_paragraph("Non-functional Requirements:")
            for item in non_functional:
                requirement = item if isinstance(item, Mapping) else {"requirement": str(item)}
                attribute = requirement.get("attribute", "Attribute")
                req_text = requirement.get("requirement") or str(item)
                rationale = requirement.get("rationale")
                document.add_paragraph(f"• {attribute}: {req_text}")
                if rationale:
                    document.add_paragraph(f"  Rationale: {rationale}")
        self._write_simple_list(document, session.get_state("dependencies"), "Dependencies")
        self._write_simple_list(document, session.get_state("risks"), "Risks")

    def _write_solution_design(self, document: Document, session: Session) -> None:
        document.add_heading("Solution Design", level=1)
        blueprint = session.get_state("solution_blueprint") or "Solution blueprint not documented."
        document.add_paragraph(blueprint)
        components = session.get_state("architecture_components") or []
        if components:
            document.add_paragraph("Component Breakdown:")
            for component in components:
                component = component if isinstance(component, Mapping) else {}
                name = component.get("name", "Component")
                responsibility = component.get("responsibility", "")
                tech_stack = component.get("tech_stack", "")
                integrations = component.get("integrations", "")
                document.add_paragraph(f"• {name} — {responsibility}")
                if tech_stack:
                    document.add_paragraph(f"  Tech: {tech_stack}")
                if integrations:
                    document.add_paragraph(f"  Integrations: {integrations}")
        self._write_simple_list(document, session.get_state("design_patterns"), "Design Patterns")
        technology_choices = session.get_state("technology_choices") or []
        if technology_choices:
            document.add_paragraph("Technology Choices:")
            for choice in technology_choices:
                choice = choice if isinstance(choice, Mapping) else {}
                area = choice.get("area", "Area")
                option = choice.get("option", "Option")
                rationale = choice.get("rationale", "")
                document.add_paragraph(f"• {area}: {option}")
                if rationale:
                    document.add_paragraph(f"  Rationale: {rationale}")
        self._write_simple_list(document, session.get_state("diagram_ideas"), "Diagram Ideas")

    def _write_prototype_development(self, document: Document, session: Session) -> None:
        document.add_heading("Prototype Development", level=1)
        self._write_simple_list(document, session.get_state("mvp_scope"), "MVP Scope")
        plan = session.get_state("prototype_plan") or []
        if plan:
            document.add_paragraph("Implementation Steps:")
            self._add_numbered_list(document, plan, prefix="Step")
        code_suggestions = session.get_state("prototype_code_suggestions") or []
        if code_suggestions:
            document.add_paragraph("Code Suggestions:")
            for suggestion in code_suggestions:
                suggestion = suggestion if isinstance(suggestion, Mapping) else {}
                desc = suggestion.get("description", "Suggestion")
                snippet = suggestion.get("snippet", "")
                document.add_paragraph(f"• {desc}")
                if snippet:
                    document.add_paragraph(snippet)
        self._write_simple_list(document, session.get_state("tooling_recommendations"), "Tooling Recommendations")
        self._write_simple_list(document, session.get_state("prototype_risks"), "Prototype Risks")
        self._write_simple_list(document, session.get_state("prototype_success_metrics"), "Prototype Success Metrics")

    def _write_testing_validation(self, document: Document, session: Session) -> None:
        document.add_heading("Testing & Validation", level=1)
        test_matrix = session.get_state("test_matrix") or []
        if test_matrix:
            document.add_paragraph("Test Matrix:")
            for entry in test_matrix:
                entry = entry if isinstance(entry, Mapping) else {}
                area = entry.get("area", "Area")
                objective = entry.get("objective", "Objective")
                test_cases = entry.get("test_cases", [])
                owner = entry.get("owner", "Owner")
                document.add_paragraph(f"• {area} — {objective} (Owner: {owner})")
                if test_cases:
                    self._add_numbered_list(document, test_cases, prefix="Test")
        self._write_simple_list(document, session.get_state("validation_plan"), "Validation Plan")
        self._write_simple_list(document, session.get_state("quality_gates"), "Quality Gates")
        self._write_simple_list(document, session.get_state("qa_tooling"), "QA Tooling")
        self._write_simple_list(document, session.get_state("qa_risks"), "QA Risks")

    def _write_documentation_summary(self, document: Document, session: Session) -> None:
        document.add_heading("Documentation Summary", level=1)
        outline = session.get_state("documentation_outline") or []
        if outline:
            for section in outline:
                section = section if isinstance(section, Mapping) else {}
                name = section.get("section", "Section")
                intent = section.get("intent", "")
                content_summary = section.get("content_summary", section.get("summary", ""))
                document.add_paragraph(f"• {name} — {intent}")
                if content_summary:
                    document.add_paragraph(f"  Summary: {content_summary}")
        self._write_simple_list(document, session.get_state("decision_log"), "Decision Log")
        self._write_simple_list(document, session.get_state("documentation_actions"), "Action Items")
        self._write_simple_list(document, session.get_state("publishing_assets"), "Publishing Assets")

    def _write_market_analysis(self, document: Document, session: Session) -> None:
        document.add_heading("Market Analysis", level=1)
        uvp = session.get_state("uvp")
        competitive = session.get_state("competitive_landscape") or []
        if uvp:
            document.add_paragraph(f"Unique Value Proposition: {uvp}")
        if competitive:
            for competitor in competitive:
                competitor = competitor if isinstance(competitor, Mapping) else {}
                name = competitor.get("name", "Competitor")
                positioning = competitor.get("positioning", "")
                strengths = competitor.get("strengths", "")
                gaps = competitor.get("gaps", "")
                document.add_paragraph(f"• {name} — Positioning: {positioning}")
                if strengths:
                    document.add_paragraph(f"  Strengths: {strengths}")
                if gaps:
                    document.add_paragraph(f"  Gaps: {gaps}")
        target_segments = session.get_state("target_segments") or []
        if target_segments:
            document.add_paragraph("Target Segments:")
            for segment in target_segments:
                segment = segment if isinstance(segment, Mapping) else {}
                name = segment.get("segment", "Segment")
                needs = segment.get("needs", "")
                fit_score = segment.get("fit_score")
                document.add_paragraph(f"• {name} — Needs: {needs} (Fit: {fit_score})")
        self._write_simple_list(document, session.get_state("go_to_market_ideas"), "Go-to-Market Ideas")
        self._write_simple_list(document, session.get_state("impact_considerations"), "Impact Considerations")

    def _write_attachments(self, document: Document, session: Session) -> None:
        document.add_heading("Attached Documents", level=1)
        attachments = session.list_attachments()
        if attachments:
            document.add_paragraph(
                f"Current strategies — Chunking: {session.chunking_strategy}, Indexing: {session.indexing_strategy}"
            )
            for attachment in attachments:
                document.add_paragraph(
                    f"• {attachment.filename} — {attachment.word_count} words, {attachment.size} bytes"
                )
        else:
            document.add_paragraph("No supporting documents attached.")

    def _write_conversation_log(self, document: Document, session: Session) -> None:
        document.add_heading("Conversation Log", level=1)
        messages = session.memory.as_list()
        for message in messages:
            role = message.get("feature") or message.get("role")
            content = message.get("content", "")
            para = add_paragraph(f"[{role}] {content}")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_numbered_list(
        self,
    document,
        items: Iterable[str],
        *,
        prefix: str,
    ) -> None:
        for idx, item in enumerate(items, start=1):
            document.add_paragraph(f"{prefix} {idx}: {item}")

    def _write_simple_list(self, document, items, heading: str) -> None:
        if not items:
            return
        add_paragraph = document.add_paragraph
        add_paragraph(f"{heading}:")
        if isinstance(items, (list, tuple, set)):
            for entry in items:
                add_paragraph(f"• {entry}")
        else:
            add_paragraph(f"• {items}")
