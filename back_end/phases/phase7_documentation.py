"""
Phase 7: Documentation

This phase generates different types of documentation based on user selection:
- Academic Report
- Software Engineering (SRS/API documentation)
- Business Proposal
- Startup Pitch
"""
from __future__ import annotations

import json
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field

from AI.llm import LLMPrompt
from common.prompts import (PHASE7_ACADEMIC_REPORT_PROMPT,
                            PHASE7_BUSINESS_PROPOSAL_PROMPT,
                            PHASE7_SRS_DOCUMENT_PROMPT,
                            PHASE7_STARTUP_PITCH_PROMPT)

# ============================================================================
# Enums
# ============================================================================

class DocumentationType(str, Enum):
    ACADEMIC_REPORT = "academic-report"
    SOFTWARE_ENGINEERING = "software-engineering"  # SRS, API docs
    BUSINESS_PROPOSAL = "business-proposal"
    STARTUP_PITCH = "startup-pitch"


# ============================================================================
# Input Models
# ============================================================================

class DocumentationInput(BaseModel):
    """Input schema for Phase 7: Documentation."""
    
    # Required
    document_type: DocumentationType = Field(
        ...,
        description="Type of documentation to generate",
    )
    
    # Optional customization
    project_title: Optional[str] = Field(
        default=None,
        description="Project title for the document",
    )
    author_name: Optional[str] = Field(
        default=None,
        description="Author or team name",
    )
    
    # User folder for export
    user_id: Optional[str] = Field(
        default=None,
        description="User ID for saving exported documents",
    )
    
    # Context from previous phases (auto-populated)
    phase1_output: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Phase 1 Problem Definition output",
    )
    phase2_output: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Phase 2 Requirements Analysis output",
    )
    phase3_output: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Phase 3 Market Analysis output",
    )


# ============================================================================
# Output Models - Academic Report
# ============================================================================

class AcademicChapter(BaseModel):
    """Chapter content for academic report."""
    chapter_number: int
    title: str
    content: str
    subsections: List[Dict[str, str]] = Field(default_factory=list)


class AcademicReportOutput(BaseModel):
    """Academic report structure following VinUni template."""
    title: str
    authors: List[str] = Field(default_factory=list)
    abstract: str
    acknowledgements: Optional[str] = None
    chapters: List[AcademicChapter]
    references: List[str]
    appendices: List[Dict[str, str]] = Field(default_factory=list)
    
    # LaTeX-specific content (aligned with VinUni template)
    latex_content: Optional[Dict[str, str]] = Field(
        default=None,
        description="LaTeX source for each section",
    )


# ============================================================================
# Output Models - SRS Document
# ============================================================================

class SrsRequirement(BaseModel):
    """Requirement in SRS format."""
    id: str
    description: str
    priority: str
    source: str
    verification: str


class SrsSection(BaseModel):
    """Section in SRS document."""
    section_number: str
    title: str
    content: str
    subsections: List[Dict[str, Any]] = Field(default_factory=list)


class SrsDocumentOutput(BaseModel):
    """SRS document following IEEE template."""
    document_title: str
    version: str
    date: str
    authors: List[str]
    
    # SRS Sections
    introduction: SrsSection
    overall_description: SrsSection
    specific_requirements: SrsSection
    external_interfaces: SrsSection
    system_features: List[Dict[str, Any]]
    non_functional_requirements: SrsSection
    
    # Additional IEEE sections
    appendices: List[Dict[str, str]] = Field(default_factory=list)
    glossary: Dict[str, str] = Field(default_factory=dict)
    
    # Raw content for template filling
    template_content: Optional[Dict[str, str]] = Field(
        default=None,
        description="Content mapped to template sections",
    )


# ============================================================================
# Output Models - Business Proposal
# ============================================================================

class BusinessProposalSection(BaseModel):
    """Section in business proposal."""
    title: str
    content: str
    key_points: List[str] = Field(default_factory=list)


class FinancialProjection(BaseModel):
    """Financial projection data."""
    metric: str
    year1: str
    year2: str
    year3: str
    assumptions: str


class BusinessProposalOutput(BaseModel):
    """Business proposal structure."""
    title: str
    executive_summary: str
    
    # Main sections
    problem_statement: BusinessProposalSection
    proposed_solution: BusinessProposalSection
    market_opportunity: BusinessProposalSection
    competitive_advantage: BusinessProposalSection
    business_model: BusinessProposalSection
    go_to_market_strategy: BusinessProposalSection
    team: BusinessProposalSection
    financial_projections: List[FinancialProjection]
    funding_request: Optional[BusinessProposalSection] = None
    risk_analysis: BusinessProposalSection
    timeline_milestones: List[Dict[str, str]]
    
    call_to_action: str


# ============================================================================
# Output Models - Startup Pitch
# ============================================================================

class PitchSlide(BaseModel):
    """Single slide in startup pitch deck."""
    slide_number: int
    title: str
    content: str
    speaker_notes: str
    visual_suggestions: List[str] = Field(default_factory=list)


class StartupPitchOutput(BaseModel):
    """Startup pitch deck structure."""
    pitch_title: str
    tagline: str
    
    # Pitch slides (standard pitch deck structure)
    slides: List[PitchSlide]
    
    # Key messages
    elevator_pitch: str  # 30-second version
    one_liner: str
    
    # Supporting materials
    key_metrics: List[Dict[str, str]]
    faqs: List[Dict[str, str]]
    
    # Presentation tips
    presentation_tips: List[str]


# ============================================================================
# Combined Output
# ============================================================================

class DocumentationOutput(BaseModel):
    """Combined output for Phase 7."""
    document_type: DocumentationType
    
    # One of these will be populated based on document_type
    academic_report: Optional[AcademicReportOutput] = None
    srs_document: Optional[SrsDocumentOutput] = None
    business_proposal: Optional[BusinessProposalOutput] = None
    startup_pitch: Optional[StartupPitchOutput] = None
    
    # Metadata
    generation_notes: List[str] = Field(default_factory=list)
    missing_information: List[str] = Field(default_factory=list)
    download_url: Optional[str] = Field(default=None, description="URL to download the exported document")


# ============================================================================
# Phase Implementation
# ============================================================================

@dataclass
class DocumentationPhase:
    """
    Phase 7: Documentation
    
    Generates various documentation types based on accumulated project data.
    Supports templates for Academic Reports (VinUni) and SRS (IEEE).
    """
    
    llm_client: Any  # LLMClient
    templates_path: str = "back_end/templates"
    
    def run(self, input_data: DocumentationInput) -> DocumentationOutput:
        """Generate documentation based on selected type."""
        
        # Build comprehensive context from all phases
        full_context = self._build_full_context(input_data)
        
        output = DocumentationOutput(document_type=input_data.document_type)
        
        if input_data.document_type == DocumentationType.ACADEMIC_REPORT:
            output.academic_report = self._generate_academic_report(full_context, input_data)
        elif input_data.document_type == DocumentationType.SOFTWARE_ENGINEERING:
            output.srs_document = self._generate_srs_document(full_context, input_data)
        elif input_data.document_type == DocumentationType.BUSINESS_PROPOSAL:
            output.business_proposal = self._generate_business_proposal(full_context, input_data)
        elif input_data.document_type == DocumentationType.STARTUP_PITCH:
            output.startup_pitch = self._generate_startup_pitch(full_context, input_data)
        
        # Track what information was missing
        output.missing_information = self._identify_missing_info(input_data)
        
        return output
    
    def _build_full_context(self, input_data: DocumentationInput) -> str:
        """Build comprehensive context from all available phase outputs."""
        context_parts = []
        
        if input_data.project_title:
            context_parts.append(f"Project Title: {input_data.project_title}")
        
        if input_data.author_name:
            context_parts.append(f"Author/Team: {input_data.author_name}")
        
        if input_data.phase1_output:
            context_parts.append(f"=== PHASE 1: PROBLEM DEFINITION ===\n{json.dumps(input_data.phase1_output, indent=2)}")
        
        if input_data.phase2_output:
            context_parts.append(f"=== PHASE 2: REQUIREMENTS ANALYSIS ===\n{json.dumps(input_data.phase2_output, indent=2)}")
        
        if input_data.phase3_output:
            context_parts.append(f"=== PHASE 3: MARKET ANALYSIS ===\n{json.dumps(input_data.phase3_output, indent=2)}")
        
        return "\n\n".join(context_parts) if context_parts else "No project context available."
    
    def _call_llm_json(self, system_prompt: str, user_prompt: str) -> Dict[str, Any]:
        """Call LLM and parse JSON response."""
        messages = [
            LLMPrompt(role="system", content=system_prompt),
            LLMPrompt(role="user", content=user_prompt),
        ]
        response = self.llm_client.generate(
            messages,
            max_tokens=4000,  # Increased for longer document generation
            extra={"response_format": {"type": "json_object"}},
        )
        content = str(response) if response else "{}"
        
        # Debug: print response for troubleshooting
        print(f"[DEBUG] LLM Response length: {len(content)} chars")
        
        try:
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            parsed = json.loads(content.strip())
            
            # Check if response has error
            if "error" in parsed:
                print(f"[DEBUG] LLM returned error: {parsed.get('error')}")
                print(f"[DEBUG] Raw content: {parsed.get('raw', content)[:500]}")
            
            return parsed
        except json.JSONDecodeError as e:
            print(f"[DEBUG] JSON parse error: {e}")
            print(f"[DEBUG] Raw content (first 500 chars): {content[:500]}")
            return {"error": "Failed to parse LLM response", "raw": content}
    
    def _generate_academic_report(
        self,
        full_context: str,
        input_data: DocumentationInput,
    ) -> AcademicReportOutput:
        """Generate academic report following VinUni template structure."""
        
        # Reference the VinUni template structure
        template_note = """
Generate content following the VinUni CECS Capstone Project Template structure:
- Abstract
- Chapter 1: Introduction (Background, Problem Statement, Objectives, Scope, Organization)
- Chapter 2: Literature Review / Related Work
- Chapter 3: Methodology / System Design
- Chapter 4: Implementation / Results
- Chapter 5: Conclusions and Future Work
- References (IEEE format)
- Appendices
"""
        context_with_template = f"{full_context}\n\n{template_note}"
        
        result = self._call_llm_json(
            PHASE7_ACADEMIC_REPORT_PROMPT,
            context_with_template,
        )
        
        chapters = [
            AcademicChapter(
                chapter_number=c.get("chapter_number", i + 1),
                title=c.get("title", ""),
                content=c.get("content", ""),
                subsections=c.get("subsections", []),
            )
            for i, c in enumerate(result.get("chapters", []))
        ]
        
        return AcademicReportOutput(
            title=result.get("title", input_data.project_title or "Capstone Project"),
            abstract=result.get("abstract", ""),
            acknowledgements=result.get("acknowledgements"),
            chapters=chapters,
            references=result.get("references", []),
            appendices=result.get("appendices", []),
            latex_content=result.get("latex_content"),
        )
    
    def _generate_srs_document(
        self,
        full_context: str,
        input_data: DocumentationInput,
    ) -> SrsDocumentOutput:
        """Generate SRS document following IEEE template with separate LLM calls per section."""
        
        project_title = input_data.project_title or "Project"
        
        def parse_section(data: Dict[str, Any], default_title: str) -> SrsSection:
            return SrsSection(
                section_number=data.get("section_number", ""),
                title=data.get("title", default_title),
                content=data.get("content", ""),
                subsections=data.get("subsections", []),
            )
        
        # Generate each section separately to avoid request failures
        print("\n[Phase 7] Generating SRS sections separately...")
        
        # 1. Introduction Section
        print("  - Generating Introduction...")
        intro_result = self._generate_srs_section_introduction(full_context, project_title)
        
        # 2. Overall Description Section
        print("  - Generating Overall Description...")
        overall_result = self._generate_srs_section_overall_description(full_context, project_title)
        
        # 3. External Interface Requirements Section
        print("  - Generating External Interface Requirements...")
        external_result = self._generate_srs_section_external_interfaces(full_context, project_title)
        
        # 4. System Features Section
        print("  - Generating System Features...")
        features_result = self._generate_srs_section_system_features(full_context, project_title)
        
        # 5. Non-Functional Requirements Section
        print("  - Generating Non-Functional Requirements...")
        nfr_result = self._generate_srs_section_nonfunctional(full_context, project_title)
        
        # 6. Other Requirements & Appendices
        print("  - Generating Other Requirements & Appendices...")
        other_result = self._generate_srs_section_other(full_context, project_title)
        
        print("  ✓ All sections generated")
        
        from datetime import datetime
        
        return SrsDocumentOutput(
            document_title=f"Software Requirements Specification for {project_title}",
            version="1.0",
            date=datetime.now().strftime("%Y-%m-%d"),
            authors=[input_data.author_name] if input_data.author_name else ["Project Team"],
            introduction=parse_section(intro_result, "Introduction"),
            overall_description=parse_section(overall_result, "Overall Description"),
            specific_requirements=parse_section(other_result.get("specific_requirements", {}), "Specific Requirements"),
            external_interfaces=parse_section(external_result, "External Interface Requirements"),
            system_features=features_result.get("system_features", []),
            non_functional_requirements=parse_section(nfr_result, "Other Nonfunctional Requirements"),
            appendices=other_result.get("appendices", []),
            glossary=other_result.get("glossary", {}),
            template_content=None,
        )

    def _generate_srs_section_introduction(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate the Introduction section (Section 1) of the SRS."""
        prompt = f"""Based on the following project context, generate the Introduction section for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with these fields for Section 1 - Introduction:
{{
    "section_number": "1",
    "title": "Introduction",
    "content": "Brief introduction paragraph",
    "subsections": [
        {{"title": "Purpose", "content": "Identify the product and describe its scope. Be specific about what this SRS covers."}},
        {{"title": "Document Conventions", "content": "Describe document standards and conventions used."}},
        {{"title": "Intended Audience and Reading Suggestions", "content": "List target readers (developers, project managers, testers, etc.) and how they should use this document."}},
        {{"title": "Product Scope", "content": "Short description of the software, its purpose, benefits, and objectives."}},
        {{"title": "References", "content": "List relevant documents, standards, and resources."}}
    ]
}}

Be specific to the project. Do not use placeholder text like <describe...>. Write actual content."""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )

    def _generate_srs_section_overall_description(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate the Overall Description section (Section 2) of the SRS."""
        prompt = f"""Based on the following project context, generate the Overall Description section for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with these fields for Section 2 - Overall Description:
{{
    "section_number": "2",
    "title": "Overall Description",
    "content": "Overview of the product and its context",
    "subsections": [
        {{"title": "Product Perspective", "content": "Describe the product's context - is it standalone, part of a larger system, replacement for existing system?"}},
        {{"title": "Product Functions", "content": "Summarize major functions as a bullet list."}},
        {{"title": "User Classes and Characteristics", "content": "Identify user types, their frequency of use, technical expertise, and needs."}},
        {{"title": "Operating Environment", "content": "Hardware platform, OS, other software it must work with."}},
        {{"title": "Design and Implementation Constraints", "content": "Regulatory policies, hardware limitations, technology constraints."}},
        {{"title": "User Documentation", "content": "Documentation to be provided (user manuals, help, tutorials)."}},
        {{"title": "Assumptions and Dependencies", "content": "Assumed factors and external dependencies."}}
    ]
}}

Be specific to the project. Write actual content, not placeholders."""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )

    def _generate_srs_section_external_interfaces(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate the External Interface Requirements section (Section 3) of the SRS."""
        prompt = f"""Based on the following project context, generate the External Interface Requirements section for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with these fields for Section 3 - External Interface Requirements:
{{
    "section_number": "3",
    "title": "External Interface Requirements",
    "content": "This section describes external interface requirements.",
    "subsections": [
        {{"title": "User Interfaces", "content": "Describe UI characteristics, screen layouts, navigation, accessibility requirements."}},
        {{"title": "Hardware Interfaces", "content": "Describe hardware the software interacts with."}},
        {{"title": "Software Interfaces", "content": "Describe connections to other software, databases, APIs, libraries."}},
        {{"title": "Communications Interfaces", "content": "Network protocols, data formats, security requirements."}}
    ]
}}

Be specific to the project. Write actual content based on the context provided."""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )

    def _generate_srs_section_system_features(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate the System Features section (Section 4) of the SRS."""
        prompt = f"""Based on the following project context, generate the System Features section for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with system features for Section 4:
{{
    "system_features": [
        {{
            "feature_id": "SF-1",
            "name": "Feature Name",
            "description": "What the feature does",
            "priority": "High/Medium/Low",
            "stimulus_response": "User actions and system responses",
            "functional_requirements": [
                {{"id": "REQ-1", "description": "Specific requirement"}},
                {{"id": "REQ-2", "description": "Another requirement"}}
            ]
        }}
    ]
}}

Include 3-5 main system features based on the project context. Each feature should have:
- Clear name and description
- Priority level
- Stimulus/response sequences
- 2-4 functional requirements with unique IDs"""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )

    def _generate_srs_section_nonfunctional(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate the Non-Functional Requirements section (Section 5) of the SRS."""
        prompt = f"""Based on the following project context, generate the Non-Functional Requirements section for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with these fields for Section 5 - Other Nonfunctional Requirements:
{{
    "section_number": "5",
    "title": "Other Nonfunctional Requirements",
    "content": "This section specifies non-functional requirements.",
    "subsections": [
        {{"title": "Performance Requirements", "content": "- NFR-PERF-001: Description of requirement\\n- NFR-PERF-002: Another requirement"}},
        {{"title": "Safety Requirements", "content": "- NFR-SAFE-001: Description of safety requirement"}},
        {{"title": "Security Requirements", "content": "- NFR-SEC-001: Description of security requirement\\n- NFR-SEC-002: Another security requirement"}},
        {{"title": "Software Quality Attributes", "content": "- NFR-QUAL-001: Reliability requirement\\n- NFR-QUAL-002: Usability requirement"}},
        {{"title": "Business Rules", "content": "- NFR-BUS-001: Business rule requirement"}}
    ]
}}

Be specific to the project. Format each requirement as a bullet list item with a unique ID (e.g., "- NFR-PERF-001: The system shall respond to user input within 2 seconds")."""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )

    def _generate_srs_section_other(self, context: str, project_title: str) -> Dict[str, Any]:
        """Generate Other Requirements, Glossary, and Appendices (Section 6+) of the SRS."""
        prompt = f"""Based on the following project context, generate the remaining sections for an IEEE 830 SRS document.

PROJECT: {project_title}

CONTEXT:
{context}

Generate a JSON object with:
{{
    "specific_requirements": {{
        "section_number": "6",
        "title": "Other Requirements",
        "content": "Brief overview of additional requirements.",
        "subsections": [
            {{"title": "Database Requirements", "content": "- DB-001: Specific database requirement for this project"}},
            {{"title": "Legal Requirements", "content": "- LEGAL-001: Specific legal/compliance requirement"}},
            {{"title": "Reuse Objectives", "content": "Description of components designed for reuse."}}
        ]
    }},
    "glossary": {{
        "Term1": "Definition of term 1",
        "Term2": "Definition of term 2"
    }},
    "appendices": [
        {{"title": "Appendix A: Glossary", "content": "See glossary section."}},
        {{"title": "Appendix B: Analysis Models", "content": "System architecture and data flow models."}},
        {{"title": "Appendix C: To Be Determined List", "content": "No TBD items."}}
    ]
}}

IMPORTANT RULES:
1. Section 6 (Other Requirements) must contain DIFFERENT content from the Appendices
2. Appendix A should just say "See glossary section." - do NOT repeat requirements here
3. Appendix B is for analysis models/diagrams - NOT for database/legal requirements
4. Appendix C is for TBD items only
5. Include 5-10 relevant technical terms in the glossary specific to this project."""
        
        return self._call_llm_json(
            "You are a technical writer creating IEEE 830 SRS documentation. Return valid JSON only.",
            prompt,
        )
    
    def _generate_business_proposal(
        self,
        full_context: str,
        input_data: DocumentationInput,
    ) -> BusinessProposalOutput:
        """Generate business proposal document."""
        
        result = self._call_llm_json(
            PHASE7_BUSINESS_PROPOSAL_PROMPT,
            full_context,
        )
        
        def parse_section(data: Dict[str, Any], default_title: str) -> BusinessProposalSection:
            return BusinessProposalSection(
                title=data.get("title", default_title),
                content=data.get("content", ""),
                key_points=data.get("key_points", []),
            )
        
        financial_projections = [
            FinancialProjection(
                metric=f.get("metric", ""),
                year1=f.get("year1", ""),
                year2=f.get("year2", ""),
                year3=f.get("year3", ""),
                assumptions=f.get("assumptions", ""),
            )
            for f in result.get("financial_projections", [])
        ]
        
        return BusinessProposalOutput(
            title=result.get("title", input_data.project_title or "Business Proposal"),
            executive_summary=result.get("executive_summary", ""),
            problem_statement=parse_section(result.get("problem_statement", {}), "Problem Statement"),
            proposed_solution=parse_section(result.get("proposed_solution", {}), "Proposed Solution"),
            market_opportunity=parse_section(result.get("market_opportunity", {}), "Market Opportunity"),
            competitive_advantage=parse_section(result.get("competitive_advantage", {}), "Competitive Advantage"),
            business_model=parse_section(result.get("business_model", {}), "Business Model"),
            go_to_market_strategy=parse_section(result.get("go_to_market_strategy", {}), "Go-to-Market Strategy"),
            team=parse_section(result.get("team", {}), "Team"),
            financial_projections=financial_projections,
            funding_request=parse_section(result.get("funding_request", {}), "Funding Request") if result.get("funding_request") else None,
            risk_analysis=parse_section(result.get("risk_analysis", {}), "Risk Analysis"),
            timeline_milestones=result.get("timeline_milestones", []),
            call_to_action=result.get("call_to_action", ""),
        )
    
    def _generate_startup_pitch(
        self,
        full_context: str,
        input_data: DocumentationInput,
    ) -> StartupPitchOutput:
        """Generate startup pitch deck."""
        
        result = self._call_llm_json(
            PHASE7_STARTUP_PITCH_PROMPT,
            full_context,
        )
        
        slides = [
            PitchSlide(
                slide_number=s.get("slide_number", i + 1),
                title=s.get("title", ""),
                content=s.get("content", ""),
                speaker_notes=s.get("speaker_notes", ""),
                visual_suggestions=s.get("visual_suggestions", []),
            )
            for i, s in enumerate(result.get("slides", []))
        ]
        
        return StartupPitchOutput(
            pitch_title=result.get("pitch_title", input_data.project_title or "Startup Pitch"),
            tagline=result.get("tagline", ""),
            slides=slides,
            elevator_pitch=result.get("elevator_pitch", ""),
            one_liner=result.get("one_liner", ""),
            key_metrics=result.get("key_metrics", []),
            faqs=result.get("faqs", []),
            presentation_tips=result.get("presentation_tips", []),
        )
    
    def _identify_missing_info(self, input_data: DocumentationInput) -> List[str]:
        """Identify what information is missing for optimal documentation."""
        missing = []
        
        if not input_data.phase1_output:
            missing.append("Phase 1 (Problem Definition) data not available - problem context may be limited")
        if not input_data.phase2_output:
            missing.append("Phase 2 (Requirements Analysis) data not available - requirements details may be limited")
        if not input_data.phase3_output:
            missing.append("Phase 3 (Market Analysis) data not available - market context may be limited")
        if not input_data.project_title:
            missing.append("Project title not specified - using default title")
        if not input_data.author_name:
            missing.append("Author/team name not specified")
        
        return missing
    
    def export_document(
        self,
        output: DocumentationOutput,
        user_id: str,
        base_path: str = "back_end/data",
    ) -> Dict[str, str]:
        """Export generated document to user's folder as DOCX.
        
        Returns dict with paths to exported files.
        """
        from datetime import datetime
        
        user_folder = Path(base_path) / user_id
        user_folder.mkdir(parents=True, exist_ok=True)
        
        exported_files = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if output.document_type == DocumentationType.SOFTWARE_ENGINEERING and output.srs_document:
            srs = output.srs_document
            
            # Debug: Check SRS content before DOCX generation
            print(f"\n[DEBUG] SRS Content Check:")
            print(f"  Title: {srs.document_title}")
            print(f"  Intro content length: {len(srs.introduction.content)}")
            print(f"  Intro content preview: {srs.introduction.content[:100] if srs.introduction.content else '(empty)'}...")
            print(f"  Intro subsections: {len(srs.introduction.subsections)}")
            
            # Export as DOCX
            docx_path = user_folder / f"srs_{timestamp}.docx"
            self._srs_to_docx(srs, docx_path)
            exported_files["docx"] = str(docx_path)
            
            # Also export JSON for verification
            json_path = user_folder / f"srs_{timestamp}.json"
            with open(json_path, 'w') as f:
                json.dump(srs.model_dump(), f, indent=2)
            exported_files["json"] = str(json_path)
            
        elif output.document_type == DocumentationType.ACADEMIC_REPORT and output.academic_report:
            report = output.academic_report
            
            # Export as DOCX
            docx_path = user_folder / f"academic_report_{timestamp}.docx"
            self._academic_report_to_docx(report, docx_path)
            exported_files["docx"] = str(docx_path)
            
        elif output.document_type == DocumentationType.BUSINESS_PROPOSAL and output.business_proposal:
            proposal = output.business_proposal
            
            # Export as DOCX
            docx_path = user_folder / f"business_proposal_{timestamp}.docx"
            self._business_proposal_to_docx(proposal, docx_path)
            exported_files["docx"] = str(docx_path)
            
        elif output.document_type == DocumentationType.STARTUP_PITCH and output.startup_pitch:
            pitch = output.startup_pitch
            
            # Export as DOCX
            docx_path = user_folder / f"startup_pitch_{timestamp}.docx"
            self._startup_pitch_to_docx(pitch, docx_path)
            exported_files["docx"] = str(docx_path)
        
        return exported_files
    
    def _get_template_path(self, template_name: str) -> Path:
        """Get the path to a template file."""
        # Try relative to this file first
        this_file = Path(__file__).resolve()
        templates_dir = this_file.parent.parent / "templates"
        template_path = templates_dir / template_name
        if template_path.exists():
            return template_path
        # Fallback to back_end/templates
        fallback = Path("back_end/templates") / template_name
        if fallback.exists():
            return fallback
        raise FileNotFoundError(f"Template not found: {template_name}")

    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str) -> bool:
        """Replace text in a paragraph while preserving formatting."""
        if old_text not in paragraph.text:
            return False
        # For simple replacement in runs
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
        # If text spans multiple runs, replace in the full paragraph text
        inline = paragraph._element
        for child in inline.iter():
            if child.text and old_text in child.text:
                child.text = child.text.replace(old_text, new_text)
                return True
        return False

    def _replace_template_placeholder(self, doc, placeholder: str, content: str) -> None:
        """Replace a <placeholder> with content in the document."""
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Check if paragraph style is 'template' - replace entire paragraph
                if paragraph.style and paragraph.style.name == 'template':
                    paragraph.clear()
                    paragraph.add_run(content)
                else:
                    self._replace_text_in_paragraph(paragraph, placeholder, content)

    @staticmethod
    def _format_numbered_lists(text: str) -> str:
        """Convert inline numbered lists like (1), (2), (3) to separate lines."""
        import re

        # Pattern to match (1), (2), etc. followed by content
        # Split on patterns like "; (2)" or ", (2)" or just "(2)"
        pattern = r'[;,]?\s*\((\d+)\)\s*'
        
        # Check if the text contains numbered items
        if not re.search(pattern, text):
            return text
        
        # Split the text by numbered markers
        parts = re.split(pattern, text)
        
        if len(parts) <= 1:
            return text
        
        # parts will be: [intro text, '1', content1, '2', content2, ...]
        result_parts = []
        
        # First part is intro text before (1)
        intro = parts[0].strip()
        if intro:
            # Remove trailing colon or semicolon
            intro = intro.rstrip(':;,').strip()
            if intro:
                result_parts.append(intro)
        
        # Process numbered items
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                num = parts[i]
                content = parts[i + 1].strip().rstrip(';,.')
                if content:
                    result_parts.append(f"({num}) {content}")
        
        return "\n".join(result_parts)
    
    @staticmethod
    def _clean_template_text(text: str) -> str:
        """Remove template placeholder text in angle brackets."""
        import re

        # Remove text within angle brackets
        cleaned = re.sub(r'<[^>]+>', '', text)
        # Clean up extra whitespace
        cleaned = re.sub(r'\n\s*\n', '\n', cleaned)
        return cleaned.strip()
    
    @staticmethod
    def _safe_str(value) -> str:
        """Convert any value to string safely."""
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            # If it's a dict, try to get 'content' key or convert to string
            return str(value.get('content', '')) if 'content' in value else str(value)
        if isinstance(value, list):
            return ", ".join(DocumentationPhase._safe_str(v) for v in value)
        return str(value)

    def _srs_to_docx(self, srs: SrsDocumentOutput, output_path: Path) -> None:
        """Fill in the IEEE SRS template with generated content."""
        import re

        from docx import Document

        # Use class-level helper methods
        safe_str = self._safe_str
        format_numbered_lists = self._format_numbered_lists
        clean_template_text = self._clean_template_text

        def get_subsection_content(sub) -> str:
            """Extract content from a subsection (could be dict or other)."""
            if isinstance(sub, dict):
                return safe_str(sub.get("content", ""))
            return safe_str(sub)

        # Load the IEEE SRS template
        template_path = self._get_template_path("srs_template-ieee.docx")
        doc = Document(template_path)
        
        # Build a mapping of section content from SRS data
        # Map subsection titles to content (ensuring string values)
        intro_subs = {safe_str(sub.get("title", "")).lower(): get_subsection_content(sub) for sub in srs.introduction.subsections}
        overall_subs = {safe_str(sub.get("title", "")).lower(): get_subsection_content(sub) for sub in srs.overall_description.subsections}
        external_subs = {safe_str(sub.get("title", "")).lower(): get_subsection_content(sub) for sub in srs.external_interfaces.subsections}
        nfr_subs = {safe_str(sub.get("title", "")).lower(): get_subsection_content(sub) for sub in srs.non_functional_requirements.subsections}
        
        # Iterate through paragraphs and fill in template placeholders
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""
            
            # Replace title page placeholders
            if text == "<Project>":
                paragraph.clear()
                paragraph.add_run(srs.document_title.replace("Software Requirements Specification for ", "").replace(" SRS", ""))
            elif "<author>" in text:
                paragraph.clear()
                authors_str = ", ".join(srs.authors) if srs.authors else "Project Team"
                paragraph.add_run(f"Prepared by {authors_str}")
            elif text == "<organization>":
                paragraph.clear()
                paragraph.add_run("VinUni CECS")
            elif text == "<date created>":
                paragraph.clear()
                paragraph.add_run(srs.date or "2024")
            
            # Replace template-style paragraphs (instruction text)
            elif style_name == "template":
                # Determine which section this template paragraph belongs to by looking at preceding headings
                section_context = self._get_section_context(doc, i)
                
                # Skip appendix paragraphs - they will be handled by _cleanup_template_placeholders
                if any(x in section_context for x in ["appendix", "glossary", "analysis models", "to be determined"]):
                    continue
                
                replacement = self._get_srs_content_for_section(
                    section_context, srs, intro_subs, overall_subs, external_subs, nfr_subs
                )
                if replacement:
                    # Clean any remaining template text and format lists
                    replacement = clean_template_text(replacement)
                    replacement = format_numbered_lists(replacement)
                    paragraph.clear()
                    paragraph.add_run(replacement)
                # Don't clear if no replacement - let cleanup handle it

        # Handle system features - find the System Feature sections and fill them
        self._fill_system_features(doc, srs.system_features)
        
        # Clean up remaining template text and appendices (handles glossary too)
        self._cleanup_template_placeholders(doc, srs)
        
        doc.save(output_path)

    def _get_section_context(self, doc, paragraph_index: int) -> str:
        """Look backwards to find the most recent heading to determine section context."""
        for i in range(paragraph_index - 1, -1, -1):
            p = doc.paragraphs[i]
            style = p.style.name if p.style else ""
            # Check for Heading styles OR TOCEntry style (used for Appendix headings)
            if style.startswith("Heading") or style == "TOCEntry":
                return p.text.strip().lower()
        return ""

    def _get_srs_content_for_section(
        self, 
        section_context: str, 
        srs: SrsDocumentOutput,
        intro_subs: Dict[str, str],
        overall_subs: Dict[str, str],
        external_subs: Dict[str, str],
        nfr_subs: Dict[str, str],
    ) -> str:
        """Get the appropriate SRS content based on section context."""
        import ast
        import json
        import re
        
        def safe_str(value) -> str:
            """Ensure value is a string, formatting lists and dicts properly."""
            if value is None:
                return ""
            
            # If string, check if it looks like JSON/Python literal and parse it
            if isinstance(value, str):
                stripped = value.strip()
                # Check if it looks like a list of dicts or a dict
                if (stripped.startswith('[{') or stripped.startswith("[{'") or 
                    stripped.startswith("{'user_class'") or stripped.startswith('{"user_class"')):
                    try:
                        # Try to parse as Python literal (handles single quotes)
                        parsed = ast.literal_eval(stripped)
                        return safe_str(parsed)
                    except (ValueError, SyntaxError):
                        pass
                    try:
                        # Try JSON parsing
                        parsed = json.loads(stripped)
                        return safe_str(parsed)
                    except json.JSONDecodeError:
                        pass
                # Remove any remaining dict-like strings that look like raw JSON
                # Pattern: {'key': 'value', ...} or {"key": "value", ...}
                if re.search(r"\{['\"]\w+['\"]:\s*['\"][^}]+['\"]\}", stripped):
                    # Attempt cleanup: extract meaningful parts
                    matches = re.findall(r"['\"]user_class['\"]:\s*['\"]([^'\"]+)['\"].*?['\"]characteristics['\"]:\s*['\"]([^'\"]+)['\"]", stripped)
                    if matches:
                        formatted = []
                        for user_class, chars in matches:
                            formatted.append(f"• {user_class}: {chars}")
                        return "\n".join(formatted)
                return value
            
            if isinstance(value, list):
                # Format list items properly
                formatted_items = []
                for item in value:
                    if isinstance(item, dict):
                        # Format dict as readable text
                        if 'user_class' in item and 'characteristics' in item:
                            # User class format
                            formatted_items.append(f"• {item['user_class']}: {item['characteristics']}")
                        elif 'id' in item and 'description' in item:
                            # Requirement format
                            formatted_items.append(f"• {item['id']}: {item['description']}")
                        elif 'name' in item:
                            formatted_items.append(f"• {item.get('name', '')}: {item.get('description', '')}")
                        else:
                            # Generic dict - format as key: value pairs
                            parts = [f"{k}: {v}" for k, v in item.items() if v]
                            formatted_items.append(f"• {'; '.join(parts)}")
                    else:
                        formatted_items.append(f"• {item}")
                return "\n".join(formatted_items)
            if isinstance(value, dict):
                if 'content' in value:
                    return safe_str(value['content'])
                # Check for user_class format in dict
                if 'user_class' in value and 'characteristics' in value:
                    return f"• {value['user_class']}: {value['characteristics']}"
                # Format dict as readable text
                parts = [f"{k}: {v}" for k, v in value.items() if v]
                return "; ".join(parts)
            return str(value)
        
        context = section_context.lower()
        
        # Helper to find matching subsection content
        def find_subsection(subs_dict: Dict[str, str], *keys: str) -> str:
            """Find content from subsection dict trying multiple possible keys."""
            for key in keys:
                val = subs_dict.get(key, "")
                if val and val.strip():
                    return safe_str(val)
            return ""
        
        # Introduction subsections
        if "purpose" in context:
            result = find_subsection(intro_subs, "purpose")
            return result or safe_str(srs.introduction.content) or f"This document specifies the software requirements for the {srs.document_title.replace('Software Requirements Specification for ', '')}."
        elif "document conventions" in context:
            result = find_subsection(intro_subs, "document conventions", "conventions")
            return result or "This document follows IEEE 830 standard conventions. Priority levels are defined as: High (must have), Medium (should have), Low (nice to have). Requirements are uniquely identified with prefixes FR (Functional), NFR (Non-Functional), and UI (User Interface)."
        elif "intended audience" in context or "reading suggestions" in context:
            result = find_subsection(intro_subs, "intended audience and reading suggestions", "intended audience", "audience")
            return result or "This document is intended for:\n• Developers: Focus on Sections 3-5 for implementation details\n• Project Managers: Review Sections 1-2 for scope and timeline implications\n• QA Teams: Concentrate on Section 5 for testing criteria\n• Stakeholders: Read Sections 1-2 for high-level understanding"
        elif "product scope" in context or context == "scope":
            result = find_subsection(intro_subs, "product scope", "scope")
            return result or safe_str(srs.introduction.content) or "This product provides comprehensive functionality as described in the system features section."
        elif "references" in context:
            result = find_subsection(intro_subs, "references")
            return result or "• IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications\n• IEEE Std 1016-2009, Software Design Descriptions\n• Project documentation and design artifacts"
        
        # Overall Description subsections
        elif "product perspective" in context:
            result = find_subsection(overall_subs, "product perspective", "perspective")
            return result or safe_str(srs.overall_description.content) or "This is a standalone software product designed to operate independently while potentially integrating with external services."
        elif "product functions" in context:
            result = find_subsection(overall_subs, "product functions", "functions")
            return result or "The major functions of this system are described in the System Features section (Section 4)."
        elif "user classes" in context or "characteristics" in context:
            result = find_subsection(overall_subs, "user classes and characteristics", "user characteristics", "user classes")
            return result or "• Primary Users: End users who will directly interact with the system\n• Administrators: Users responsible for system configuration and maintenance\n• Developers: Technical users who may extend or integrate with the system"
        elif "operating environment" in context:
            result = find_subsection(overall_subs, "operating environment", "environment")
            return result or "The system is designed to operate on standard computing platforms with modern web browsers or mobile devices."
        elif "design and implementation" in context or "constraints" in context:
            result = find_subsection(overall_subs, "design and implementation constraints", "constraints", "design constraints")
            return result or "• Must comply with relevant data protection regulations\n• Must be accessible on standard hardware\n• Must follow secure coding practices"
        elif "user documentation" in context:
            result = find_subsection(overall_subs, "user documentation", "documentation")
            return result or "User documentation will include:\n• User Guide: Step-by-step instructions for all features\n• Quick Start Guide: Getting started documentation\n• Online Help: Context-sensitive help within the application"
        elif "assumptions" in context or "dependencies" in context:
            result = find_subsection(overall_subs, "assumptions and dependencies", "assumptions", "dependencies")
            return result or "• Users have access to compatible devices\n• Network connectivity is available for online features\n• External services maintain their documented APIs"
        
        # External Interface Requirements
        elif "user interfaces" in context:
            result = find_subsection(external_subs, "user interfaces", "ui")
            return result or safe_str(srs.external_interfaces.content) or "The user interface shall be intuitive and accessible, following platform-specific design guidelines."
        elif "hardware interfaces" in context:
            result = find_subsection(external_subs, "hardware interfaces", "hardware")
            return result or "The software interfaces with standard computing hardware including processors, memory, storage, and display devices."
        elif "software interfaces" in context:
            result = find_subsection(external_subs, "software interfaces", "software")
            return result or "The system interfaces with the operating system, relevant APIs, and any external services required for functionality."
        elif "communications interfaces" in context:
            result = find_subsection(external_subs, "communications interfaces", "communications", "network")
            return result or "The system uses standard network protocols (HTTP/HTTPS) for any network communications."
        
        # System Features
        elif "system feature" in context:
            return ""  # Handled separately
        
        # Non-functional Requirements
        elif "performance" in context:
            result = find_subsection(nfr_subs, "performance requirements", "performance")
            return result or "• Response Time: The system shall respond to user input within 2 seconds under normal load\n• Throughput: The system shall support the expected number of concurrent users\n• Resource Utilization: The system shall operate efficiently within available resources"
        elif "safety" in context:
            result = find_subsection(nfr_subs, "safety requirements", "safety")
            return result or "• Data Integrity: The system shall prevent data corruption through validation and error handling\n• Graceful Degradation: The system shall handle failures gracefully without data loss"
        elif "security" in context:
            result = find_subsection(nfr_subs, "security requirements", "security")
            return result or "• Authentication: The system shall implement secure user authentication\n• Authorization: Access control shall be enforced based on user roles\n• Data Protection: Sensitive data shall be encrypted in transit and at rest"
        elif "software quality" in context:
            result = find_subsection(nfr_subs, "software quality attributes", "quality", "quality attributes")
            return result or "• Reliability: The system shall maintain high availability\n• Maintainability: Code shall follow established standards for ease of maintenance\n• Usability: The interface shall be intuitive for target users"
        elif "business rules" in context:
            result = find_subsection(nfr_subs, "business rules", "rules")
            return result or "Business rules governing system behavior are defined in conjunction with stakeholder requirements."
        
        # Other Requirements (Section 6)
        elif "other requirements" in context:
            if srs.specific_requirements:
                # Build content from specific_requirements section and its subsections
                parts = []
                main_content = safe_str(srs.specific_requirements.content)
                if main_content:
                    parts.append(main_content)
                # Add subsection content
                for sub in srs.specific_requirements.subsections:
                    if isinstance(sub, dict):
                        sub_title = sub.get("title", "")
                        sub_content = sub.get("content", "")
                        if sub_title and sub_content:
                            parts.append(f"\n{sub_title}:\n{sub_content}")
                if parts:
                    return "\n".join(parts)
            # Default fallback for Other Requirements
            return "This section covers additional requirements not addressed elsewhere:\n• Database Requirements: System data storage and retrieval specifications\n• Legal and Regulatory: Compliance with applicable laws and standards\n• Reuse Objectives: Components designed for potential reuse in future projects"
        elif "database" in context:
            # Handle database requirements subsection
            if srs.specific_requirements and srs.specific_requirements.subsections:
                for sub in srs.specific_requirements.subsections:
                    if isinstance(sub, dict) and "database" in sub.get("title", "").lower():
                        return safe_str(sub.get("content", ""))
            return ""
        elif "legal" in context:
            # Handle legal requirements subsection
            if srs.specific_requirements and srs.specific_requirements.subsections:
                for sub in srs.specific_requirements.subsections:
                    if isinstance(sub, dict) and "legal" in sub.get("title", "").lower():
                        return safe_str(sub.get("content", ""))
            return ""
        elif "glossary" in context:
            return ""  # Handled separately
        elif "analysis models" in context:
            return "See appendices for detailed analysis models."
        elif "to be determined" in context:
            return "No items are currently marked as TBD."
        
        return ""

    def _fill_system_features(self, doc, system_features: List[Dict[str, Any]]) -> None:
        """Fill in system features section of the template."""
        import re

        from docx.shared import Pt
        
        if not system_features:
            return
        
        # Find paragraphs to clear (template placeholders in System Features section)
        paragraphs_to_clear = []
        feature_1_index = -1
        in_system_features_section = False
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            text_lower = text.lower()
            style_name = p.style.name if p.style else ""
            
            # Track when we enter System Features section (Section 4)
            if "system features" in text_lower and style_name == "Heading 1":
                in_system_features_section = True
            
            # Track when we exit (Section 5 or Section 6)
            if style_name == "Heading 1" and ("nonfunctional" in text_lower or "other requirements" in text_lower):
                in_system_features_section = False
            
            if in_system_features_section:
                if "system feature 1" in text_lower and style_name == "Heading 2":
                    feature_1_index = i
                # Mark template sub-headings and placeholders for clearing
                elif "system feature 2" in text_lower or "(and so on)" in text_lower:
                    paragraphs_to_clear.append(i)
                # Clear ALL Heading 2 templates like "4.2 System Feature 2", "4.3", etc.
                elif style_name == "Heading 2" and re.match(r'4\.\d+', text):
                    paragraphs_to_clear.append(i)
                # Clear "level 4" sub-headings (4.1.1, 4.1.2, 4.1.3)
                elif style_name == "level 4":
                    paragraphs_to_clear.append(i)
                # Also clear Heading 3 just in case
                elif style_name == "Heading 3":
                    paragraphs_to_clear.append(i)
                # Clear "level 3 text" template content
                elif style_name == "level 3 text":
                    paragraphs_to_clear.append(i)
                # Clear template style paragraphs
                elif style_name == "template" or re.search(r'<[^>]+>', text):
                    paragraphs_to_clear.append(i)
                # Clear requirement placeholder lines
                elif style_name == "requirement" or text.startswith("REQ-"):
                    paragraphs_to_clear.append(i)
        
        # Clear marked paragraphs
        for idx in paragraphs_to_clear:
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].clear()
        
        if feature_1_index < 0:
            return
            
        # Find where to insert features (after "System Feature 1" heading)
        insert_after_index = feature_1_index
        
        # Fill first feature in place of "System Feature 1"
        first_feature = system_features[0]
        p = doc.paragraphs[feature_1_index]
        p.clear()
        run = p.add_run(f"{first_feature.get('name', 'Feature 1')}")
        
        # Find and fill the description paragraph after the heading
        for j in range(feature_1_index + 1, min(feature_1_index + 15, len(doc.paragraphs))):
            next_p = doc.paragraphs[j]
            style_name = next_p.style.name if next_p.style else ""
            
            # Skip heading 3s and find template/content paragraphs
            if style_name in ["template", "level 3 text", "Normal"]:
                next_p.clear()
                
                desc = first_feature.get('description', 'No description provided.')
                priority = first_feature.get('priority', 'Medium')
                stimulus = first_feature.get('stimulus_response', first_feature.get('stimulus', ''))
                func_reqs = first_feature.get('functional_requirements', [])
                
                # Build feature content
                content_parts = []
                content_parts.append(f"Description: {desc}")
                content_parts.append(f"Priority: {priority}")
                if stimulus:
                    content_parts.append(f"\nStimulus/Response: {stimulus}")
                if func_reqs:
                    content_parts.append("\nFunctional Requirements:")
                    if isinstance(func_reqs, list):
                        for req in func_reqs:
                            if isinstance(req, dict):
                                req_id = req.get('id', req.get('name', ''))
                                req_desc = req.get('description', req.get('content', ''))
                                content_parts.append(f"  • {req_id}: {req_desc}")
                            else:
                                content_parts.append(f"  • {req}")
                    else:
                        content_parts.append(f"  {func_reqs}")
                
                next_p.add_run("\n".join(content_parts))
                insert_after_index = j
                break
        
        # Add additional features after the first one
        if len(system_features) > 1:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Find the paragraph after which to insert
            insert_element = doc.paragraphs[insert_after_index]._element
            
            for idx, feature in enumerate(system_features[1:], start=2):
                # Add a blank line
                blank_p = doc.add_paragraph()
                insert_element.addnext(blank_p._element)
                insert_element = blank_p._element
                
                # Add heading for this feature (without 4.X prefix - template already has numbering)
                heading_p = doc.add_paragraph()
                heading_p.style = doc.styles['Heading 2']
                heading_p.add_run(f"{feature.get('name', f'Feature {idx}')}")
                insert_element.addnext(heading_p._element)
                insert_element = heading_p._element
                
                # Add feature content
                content_p = doc.add_paragraph()
                
                desc = feature.get('description', 'No description provided.')
                priority = feature.get('priority', 'Medium')
                stimulus = feature.get('stimulus_response', feature.get('stimulus', ''))
                func_reqs = feature.get('functional_requirements', [])
                
                content_parts = []
                content_parts.append(f"Description: {desc}")
                content_parts.append(f"Priority: {priority}")
                if stimulus:
                    content_parts.append(f"\nStimulus/Response: {stimulus}")
                if func_reqs:
                    content_parts.append("\nFunctional Requirements:")
                    if isinstance(func_reqs, list):
                        for req in func_reqs:
                            if isinstance(req, dict):
                                req_id = req.get('id', req.get('name', ''))
                                req_desc = req.get('description', req.get('content', ''))
                                content_parts.append(f"  • {req_id}: {req_desc}")
                            else:
                                content_parts.append(f"  • {req}")
                    else:
                        content_parts.append(f"  {func_reqs}")
                
                content_p.add_run("\n".join(content_parts))
                insert_element.addnext(content_p._element)
                insert_element = content_p._element

    def _fill_glossary(self, doc, glossary: Dict[str, str]) -> None:
        """Fill in glossary section of the template."""
        if not glossary:
            return
            
        for i, p in enumerate(doc.paragraphs):
            if "appendix a: glossary" in p.text.lower() or (p.style and "glossary" in p.text.lower()):
                # Find the next template paragraph and fill with glossary
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_p = doc.paragraphs[j]
                    if next_p.style and next_p.style.name == "template":
                        glossary_text = "\n".join([f"{term}: {defn}" for term, defn in glossary.items()])
                        next_p.clear()
                        next_p.add_run(glossary_text)
                        break
                break
    
    def _cleanup_template_placeholders(self, doc, srs: 'SrsDocumentOutput') -> None:
        """Remove all remaining template placeholder text and clean up appendices."""
        import re

        # Track which section/appendix we're in
        current_section = None
        appendix_filled = {"glossary": False, "analysis": False, "tbd": False}
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            style_name = p.style.name if p.style else ""
            text_lower = text.lower()
            
            # Track when we enter/exit Section 6 (Other Requirements) or appendices
            # Appendices use "TOCEntry" style in this template
            if style_name == "Heading 1":
                if "other requirements" in text_lower:
                    current_section = "section6"
            elif style_name == "TOCEntry":
                if "appendix a" in text_lower or "glossary" in text_lower:
                    current_section = "glossary"
                elif "appendix b" in text_lower or "analysis" in text_lower:
                    current_section = "analysis"
                elif "appendix c" in text_lower or "to be determined" in text_lower:
                    current_section = "tbd"
            
            # Only remove paragraphs that STILL contain template instructions (text in angle brackets)
            # Do NOT clear based on style alone - content may have been filled in
            has_template_text = re.search(r'<[^>]+>', text)
            
            if has_template_text:
                # Check if this is in an appendix and handle appropriately
                if current_section == "glossary" and not appendix_filled["glossary"]:
                    if srs.glossary:
                        glossary_text = "\n".join([f"• {term}: {defn}" for term, defn in srs.glossary.items()])
                        p.clear()
                        p.add_run(glossary_text)
                        appendix_filled["glossary"] = True
                    else:
                        p.clear()
                elif current_section == "analysis" and not appendix_filled["analysis"]:
                    # Fill with proper analysis models content (NOT Section 6 content)
                    p.clear()
                    p.add_run("Data flow diagrams, class diagrams, and system architecture models are documented separately.")
                    appendix_filled["analysis"] = True
                elif current_section == "tbd" and not appendix_filled["tbd"]:
                    # Fill with proper TBD content (NOT Section 6 content)
                    p.clear()
                    p.add_run("No items are currently marked as To Be Determined.")
                    appendix_filled["tbd"] = True
                else:
                    # Just clear template text that wasn't filled
                    p.clear()
            
            # Remove "level 4" sub-headings (4.1.1 Description and Priority, etc.)
            if style_name == "level 4":
                p.clear()
            
            # Remove "level 3 text" template content
            if style_name == "level 3 text":
                p.clear()
            
            # Remove template sub-headings in System Features (4.1.1, 4.1.2, 4.1.3) - also check Heading 3
            if style_name == "Heading 3":
                if any(x in text_lower for x in ["description and priority", "stimulus/response", "functional requirements"]):
                    p.clear()
            
            # Remove empty Heading 2 sections (like "4.6" with no real content)
            if style_name == "Heading 2":
                # Check if this looks like an empty numbered section or template feature
                # Match patterns like "4.6", "4.6 System Feature", etc.
                if re.match(r'^4\.\d+\s*$', text):  # Just "4.6" or "4.6 "
                    p.clear()
                elif re.match(r'^4\.\d+', text) and ("system feature" in text_lower or text_lower.strip() == text.strip().lower()):
                    p.clear()
            
            # Remove placeholder requirement IDs (requirement style or REQ- prefix)
            if style_name == "requirement" or text.startswith("REQ-"):
                p.clear()
    
    def _academic_report_to_docx(self, report: AcademicReportOutput, output_path: Path) -> None:
        """Fill in the VinUni Capstone template with generated content."""
        import re

        from docx import Document

        # Load the VinUni template
        template_path = self._get_template_path("VinUni-CECS-Capstone-Project-template.docx")
        doc = Document(template_path)
        
        # Extract author names from report or use default
        authors = report.title.split(" by ")[-1] if " by " in report.title else "Project Team"
        project_title = report.title.replace(f" by {authors}", "") if " by " in report.title else report.title
        
        # Build chapter content mapping
        chapter_map = {}
        for ch in report.chapters:
            ch_title = ch.title.lower()
            chapter_map[ch_title] = ch
        
        # Helper to format content
        def format_content(text: str) -> str:
            """Apply formatting to content text."""
            text = self._clean_template_text(text)
            text = self._format_numbered_lists(text)
            return text
        
        # Fill in template placeholders
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""
            
            # Title page replacements
            if text == "Title of Your Capstone Project":
                paragraph.clear()
                paragraph.add_run(project_title)
            elif "First Author" in text and "243014000" in text:
                # Author lines - replace with actual authors if available
                if hasattr(report, 'authors') and report.authors:
                    paragraph.clear()
                    paragraph.add_run(report.authors[0] if len(report.authors) > 0 else "Author 1")
            elif text == "Second Author\t243014000":
                paragraph.clear()
                if hasattr(report, 'authors') and len(report.authors) > 1:
                    paragraph.add_run(report.authors[1])
            elif text == "Third Author\t243014000":
                paragraph.clear()
                if hasattr(report, 'authors') and len(report.authors) > 2:
                    paragraph.add_run(report.authors[2])
            elif "Supervisor's Name" in text:
                paragraph.clear()
                paragraph.add_run("Supervisor: Faculty Advisor")
            elif text == "September, 2023":
                from datetime import datetime
                paragraph.clear()
                paragraph.add_run(datetime.now().strftime("%B, %Y"))
            
            # Abstract section
            elif text == "Write the abstract here.":
                paragraph.clear()
                paragraph.add_run(format_content(report.abstract))
            
            # Acknowledgements
            elif "I would like to extend my sincere thanks" in text:
                if report.acknowledgements:
                    paragraph.clear()
                    paragraph.add_run(format_content(report.acknowledgements))
            
            # Chapter content - look for chapter headings and fill content
            elif style_name == "Normal" and text == "INTRODUCTION":
                self._fill_chapter_content(doc, i, chapter_map.get("introduction", None))
            
            # Handle template instruction paragraphs
            elif "Chapter 1 should provide" in text or "Provide literature review" in text:
                # Find Introduction chapter and fill
                intro_ch = chapter_map.get("introduction", None)
                if intro_ch:
                    paragraph.clear()
                    paragraph.add_run(format_content(intro_ch.content))
            
            elif "Problem Statement" in text and style_name == "Heading 4":
                # Look for problem statement content in Introduction chapter
                intro_ch = chapter_map.get("introduction", None)
                if intro_ch:
                    for sub in intro_ch.subsections:
                        if "problem" in sub.get("title", "").lower():
                            # Fill next paragraph
                            self._fill_next_body_paragraph(doc, i, format_content(sub.get("content", "")))
                            break
            
            elif "State the problem in a single sentence" in text:
                intro_ch = chapter_map.get("introduction", None)
                if intro_ch:
                    paragraph.clear()
                    # Extract problem statement from subsections
                    for sub in intro_ch.subsections:
                        if "problem" in sub.get("title", "").lower():
                            paragraph.add_run(format_content(sub.get("content", "")))
                            break
            
            # Clear any remaining template placeholder text (in angle brackets)
            elif re.search(r'<[^>]+>', text):
                paragraph.clear()
        
        # Fill chapter bodies with formatting
        self._fill_report_chapters(doc, report.chapters)
        
        # Fill references
        self._fill_references(doc, report.references)
        
        # Final cleanup - remove any remaining template instruction text
        self._cleanup_academic_template(doc)
        
        doc.save(output_path)
    
    def _cleanup_academic_template(self, doc) -> None:
        """Remove remaining template instruction text from academic report."""
        import re

        # List of common template instruction phrases to remove
        template_phrases = [
            "should provide",
            "provide a brief",
            "describe the",
            "explain the",
            "outline the",
            "summarize the",
            "write the",
            "include a",
            "present the",
            "discuss the",
        ]
        
        for p in doc.paragraphs:
            text = p.text.strip().lower()
            style_name = p.style.name if p.style else ""
            
            # Clear paragraphs with template placeholders
            if re.search(r'<[^>]+>', p.text):
                p.clear()
            # Clear instruction paragraphs (but not headings)
            elif style_name not in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
                # Check for template instruction phrases in body text
                if any(phrase in text for phrase in template_phrases) and len(text) > 100:
                    # Only clear if it looks like instruction text (not actual content)
                    if "chapter" in text or "section" in text or "should" in text:
                        p.clear()

    def _fill_chapter_content(self, doc, start_idx: int, chapter) -> None:
        """Fill chapter content after a chapter heading."""
        if not chapter:
            return
        # Find next body text paragraph and fill with chapter content
        for j in range(start_idx + 1, min(start_idx + 20, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            style = p.style.name if p.style else ""
            if style in ["Body Text", "Normal"] and "should provide" in p.text.lower():
                p.clear()
                # Apply formatting
                content = self._clean_template_text(chapter.content)
                content = self._format_numbered_lists(content)
                p.add_run(content)
                break

    def _fill_next_body_paragraph(self, doc, start_idx: int, content: str) -> None:
        """Fill the next body text paragraph with content."""
        for j in range(start_idx + 1, min(start_idx + 5, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            style = p.style.name if p.style else ""
            if style in ["Body Text", "Normal"]:
                p.clear()
                p.add_run(content)
                break

    def _fill_report_chapters(self, doc, chapters: List[AcademicChapter]) -> None:
        """Fill in chapter content throughout the document."""
        # Map chapter numbers to titles in template
        template_chapters = {
            1: "INTRODUCTION",
            2: "PROJECT MANAGEMENT",  # Could be Literature Review
            3: "METHODOLOGY",
            4: "IMPLEMENTATION",
            5: "CONCLUSIONS",
        }
        
        for chapter in chapters:
            ch_num = chapter.chapter_number
            # Find the chapter heading in the document
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip() == f"Chapter {ch_num}" or f"CHAPTER {ch_num}" in p.text.upper():
                    # Fill content in subsequent paragraphs
                    for j in range(i + 1, min(i + 30, len(doc.paragraphs))):
                        next_p = doc.paragraphs[j]
                        if next_p.style and next_p.style.name in ["Body Text"] and len(next_p.text) > 50:
                            # This is instruction text, replace with content
                            next_p.clear()
                            # Apply formatting
                            content = self._clean_template_text(chapter.content)
                            content = self._format_numbered_lists(content)
                            next_p.add_run(content)
                            break
                    break

    def _fill_references(self, doc, references: List[str]) -> None:
        """Fill references section in the template."""
        if not references:
            return
            
        # Find references section
        for i, p in enumerate(doc.paragraphs):
            if "REFERENCES" in p.text.upper() or "Bibliography" in p.text:
                # Add references after this heading
                for j, ref in enumerate(references, 1):
                    # Find a place to insert or append
                    if i + j < len(doc.paragraphs):
                        ref_p = doc.paragraphs[i + j]
                        if ref_p.style and ref_p.style.name == "Body Text":
                            ref_p.clear()
                            ref_p.add_run(f"[{j}] {ref}")
                break
    
    def _business_proposal_to_docx(self, proposal: BusinessProposalOutput, output_path: Path) -> None:
        """Convert business proposal to DOCX format."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Helper to format content
        def fmt(text: str) -> str:
            text = self._clean_template_text(text)
            text = self._format_numbered_lists(text)
            return text
        
        # Title
        title = doc.add_heading(proposal.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(fmt(proposal.executive_summary))
        
        # Main sections
        sections = [
            proposal.problem_statement,
            proposal.proposed_solution,
            proposal.market_opportunity,
            proposal.competitive_advantage,
            proposal.business_model,
            proposal.go_to_market_strategy,
            proposal.team,
            proposal.risk_analysis,
        ]
        
        for section in sections:
            doc.add_heading(section.title, level=1)
            doc.add_paragraph(fmt(section.content))
            if section.key_points:
                doc.add_paragraph("Key Points:", style="Intense Quote")
                for point in section.key_points:
                    doc.add_paragraph(f"• {fmt(point)}")
        
        # Financial Projections
        if proposal.financial_projections:
            doc.add_heading("Financial Projections", level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Year 1"
            hdr_cells[2].text = "Year 2"
            hdr_cells[3].text = "Year 3"
            hdr_cells[4].text = "Assumptions"
            for proj in proposal.financial_projections:
                row_cells = table.add_row().cells
                row_cells[0].text = fmt(proj.metric)
                row_cells[1].text = fmt(proj.year1)
                row_cells[2].text = fmt(proj.year2)
                row_cells[3].text = fmt(proj.year3)
                row_cells[4].text = fmt(proj.assumptions)
        
        # Timeline & Milestones
        if proposal.timeline_milestones:
            doc.add_heading("Timeline & Milestones", level=1)
            for milestone in proposal.timeline_milestones:
                p = doc.add_paragraph()
                p.add_run(f"{milestone.get('date', '')}: ").bold = True
                p.add_run(fmt(milestone.get("milestone", "")))
        
        # Call to Action
        doc.add_heading("Call to Action", level=1)
        doc.add_paragraph(fmt(proposal.call_to_action))
        
        doc.save(output_path)
    
    def _startup_pitch_to_docx(self, pitch: StartupPitchOutput, output_path: Path) -> None:
        """Convert startup pitch to DOCX format."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Helper to format content
        def fmt(text: str) -> str:
            text = self._clean_template_text(text)
            text = self._format_numbered_lists(text)
            return text
        
        # Title
        title = doc.add_heading(pitch.pitch_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        tagline = doc.add_paragraph(fmt(pitch.tagline))
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.runs[0].italic = True
        
        doc.add_page_break()
        
        # Elevator Pitch
        doc.add_heading("Elevator Pitch (30 seconds)", level=1)
        doc.add_paragraph(fmt(pitch.elevator_pitch))
        
        # One-Liner
        doc.add_heading("One-Liner", level=1)
        doc.add_paragraph(fmt(pitch.one_liner))
        
        doc.add_page_break()
        
        # Slides
        doc.add_heading("Pitch Deck Slides", level=1)
        for slide in pitch.slides:
            doc.add_heading(f"Slide {slide.slide_number}: {slide.title}", level=2)
            doc.add_paragraph(fmt(slide.content))
            
            notes = doc.add_paragraph()
            notes.add_run("Speaker Notes: ").bold = True
            notes.add_run(fmt(slide.speaker_notes)).italic = True
            
            if slide.visual_suggestions:
                doc.add_paragraph("Visual Suggestions:")
                for vs in slide.visual_suggestions:
                    doc.add_paragraph(f"• {fmt(vs)}")
        
        # Key Metrics
        if pitch.key_metrics:
            doc.add_heading("Key Metrics", level=1)
            for metric in pitch.key_metrics:
                p = doc.add_paragraph()
                p.add_run(f"{metric.get('name', '')}: ").bold = True
                p.add_run(fmt(metric.get("value", "")))
        
        # FAQs
        if pitch.faqs:
            doc.add_heading("FAQs", level=1)
            for faq in pitch.faqs:
                q = doc.add_paragraph()
                q.add_run(f"Q: {faq.get('question', '')}").bold = True
                doc.add_paragraph(f"A: {fmt(faq.get('answer', ''))}")
        
        # Presentation Tips
        if pitch.presentation_tips:
            doc.add_heading("Presentation Tips", level=1)
            for tip in pitch.presentation_tips:
                doc.add_paragraph(f"• {fmt(tip)}")
        
        doc.save(output_path)
